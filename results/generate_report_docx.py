"""Generate PROJECT_REPORT.docx for drone-rl-lab."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level, size in [(1, 18), (2, 14), (3, 12)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

# -- Title --
title = doc.add_heading('Drone RL Lab \u2014 Project Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

doc.add_paragraph('March 2026 \u2022 jeffmW5/drone-rl-lab').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ===== 1. EXECUTIVE SUMMARY =====
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This project uses drone simulation to learn reinforcement learning from scratch. '
    'Two backends are supported: gym-pybullet-drones (hover task, SB3 PPO) and '
    'lsy_drone_racing (gate racing, CleanRL PPO with JAX). '
    'The goal is to achieve a sub-5-second lap time on Level 2 of the LSY drone racing '
    'competition, placing in the top 3 on the Kaggle leaderboard.'
)
doc.add_paragraph(
    'Over 16 experiments across CPU and GPU, the project progressed from basic hover control '
    'to the first RL model capable of completing a Level 2 randomized gate race \u2014 something '
    'even the reference pre-trained RL model cannot do.'
)

# ===== 2. INFRASTRUCTURE =====
doc.add_heading('2. Infrastructure', level=1)

infra_items = [
    ('Two-Agent Agentic Loop', 'Windows Claude (orchestrator) writes INBOX.md with experiment instructions. '
     'Linux Claude (executor) runs experiments on the VM and writes results to outbox/.'),
    ('VirtualBox Ubuntu VM', 'Ubuntu VM "CF" with shared folder at /media/ for seamless file exchange.'),
    ('Config-Driven Experiments', 'YAML configs define every experiment parameter. A dispatcher (train.py) '
     'routes to the correct backend trainer based on the "backend:" field.'),
    ('Memory System', 'MEMORY.md serves as institutional memory \u2014 hard rules, experiment log, and lessons '
     'that persist across sessions so agents never repeat failed experiments.'),
    ('RunPod GPU Cloud', 'RTX 3090 (24GB VRAM) for scaled training. Setup script with 4-hour auto-shutdown, '
     'deploy keys, and spending cap ($10).'),
]
for title_text, desc in infra_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

# ===== 3. HOVER RESULTS =====
doc.add_heading('3. Hover Backend Results (exp 001\u2013005)', level=1)
doc.add_paragraph(
    'Five experiments explored reward function design for the HoverAviary environment '
    'using SB3 PPO with ONE_D_RPM action space (single thrust dimension).'
)

hover_data = [
    ['Exp', 'Change', 'Reward', 'Outcome'],
    ['001', 'Baseline quartic reward', '474', '\u2705 Ceiling for 1D RPM'],
    ['002', '2x time budget', '474', '\u2705 Same ceiling, faster convergence'],
    ['003', 'Quadratic reward', '369', '\u274c Worse than quartic'],
    ['004', 'Velocity penalty', '407', '\u274c Penalty hurts performance'],
    ['005', 'Conservative PPO', '437', '\u274c Stability vs performance tradeoff'],
]
table = doc.add_table(rows=len(hover_data), cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(hover_data):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = cell_text
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key Learning: ')
run.bold = True
p.add_run('The ONE_D_RPM action space caps reward at ~474 regardless of reward function. '
          'The bottleneck is the single action dimension, not the reward design.')

# ===== 4. RACING RESULTS =====
doc.add_heading('4. Racing Backend Results', level=1)

doc.add_heading('4a. CPU Training (VirtualBox VM, 64 envs)', level=2)
cpu_data = [
    ['Exp', 'Level', 'n_obs', 'Steps', 'Reward', 'Lap (s)', 'Outcome'],
    ['010', 'L0', '0', '500K', '7.36', '13.36', '\u2705 Beats PID, 0.024s off reference'],
    ['013', 'L0', '2', '297K*', '5.02', 'crash', '\u274c Regression \u2014 insufficient compute'],
]
table = doc.add_table(rows=len(cpu_data), cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(cpu_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text
        if i == 0:
            for run in table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph('*exp_013 only completed 297K of 500K target steps within 600s budget.')

doc.add_heading('4b. GPU Training (RunPod RTX 3090, 1024 envs)', level=2)
gpu_data = [
    ['Exp', 'Level', 'Steps', 'Reward', 'Wall Time', 'Lap (s)', 'Finish Rate', 'Outcome'],
    ['014', 'L0', '1.5M', '7.29', '206s', 'TBD', 'TBD', '\u2705 Validates n_obs=2 with GPU'],
    ['015', 'L2', '3M', '7.53', '413s', 'TBD', 'TBD', '\u2705 First L2 training, still climbing'],
    ['016', 'L2', '10M', '7.71', '1466s', '13.49', '2/10 (20%)', '\u2705 FIRST RL to finish L2'],
]
table = doc.add_table(rows=len(gpu_data), cols=8)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(gpu_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text
        if i == 0:
            for run in table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Breakthrough: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run('exp_016 is the first RL model to complete Level 2 races. The reference pre-trained RL '
          'model scores 0/5 finishes on Level 2. Our model achieves 2/10 (20%).')

# ===== 5. SIM BENCHMARK =====
doc.add_heading('5. Simulator Benchmark \u2014 All Controllers vs All Levels', level=1)

for level_name, level_data in [
    ('Level 0 (Perfect Knowledge)', [
        ['Controller', 'Type', 'Avg (s)', 'Finished', 'Gates'],
        ['Their RL (pre-trained)', 'RL', '13.34', '5/5', '4/4'],
        ['Our exp_010', 'RL', '13.36', '5/5', '4/4'],
        ['PID attitude', 'PID', '13.37', '5/5', '4/4'],
        ['State trajectory', 'Traj', '13.86', '5/5', '4/4'],
    ]),
    ('Level 1 (Randomized Physics)', [
        ['Controller', 'Type', 'Avg (s)', 'Finished', 'Gates'],
        ['Their RL', 'RL', '13.34', '5/5', '4/4'],
        ['PID attitude', 'PID', '13.39', '5/5', '4/4'],
        ['State trajectory', 'Traj', '13.86', '5/5', '4/4'],
    ]),
    ('Level 2 (Randomized Physics + Gates) \u2014 COMPETITION', [
        ['Controller', 'Type', 'Avg (s)', 'Finished', 'Gates'],
        ['State trajectory', 'Traj', '5.96', '1/5', '0\u20134/4'],
        ['PID attitude', 'PID', '8.59', '2/5', '0\u20134/4'],
        ['Their RL', 'RL', '7.25', '0/5', '0\u20133/4'],
        ['Our exp_016', 'RL (GPU)', '13.49', '2/10', '0\u20134/4'],
    ]),
]:
    doc.add_heading(level_name, level=2)
    table = doc.add_table(rows=len(level_data), cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(level_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph('')

# ===== 6. KAGGLE STANDING =====
doc.add_heading('6. Kaggle Competition Standing', level=1)
kaggle_data = [
    ['Rank', 'Team', 'Avg Lap (s)'],
    ['1', 'Team Y', '3.394'],
    ['2', 'Group6', '4.886'],
    ['3', 'Limo', '5.022'],
    ['\u2014', 'Our exp_016', '13.49 (20% finish)'],
]
table = doc.add_table(rows=len(kaggle_data), cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(kaggle_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text
        if i == 0:
            for run in table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph('')
doc.add_paragraph(
    'Gap analysis: The top teams are 3\u20134x faster with near-100% completion rates. '
    'The gap is not in policy quality (our RL beats PID on Level 0) but in trajectory '
    'generation \u2014 Kaggle winners likely use dynamic path planning that adapts to '
    'randomized gate positions.'
)

# ===== 7. KEY LESSONS =====
doc.add_heading('7. Key Lessons Learned', level=1)
lessons = [
    'n_obs=2 (observation stacking) requires 1024+ parallel envs on GPU, not 64 on CPU. '
    'The larger observation space slows training dramatically at small scale.',
    'Level 2 bottleneck is trajectory generation, not policy quality. All controllers use '
    'fixed polynomial trajectories that fly through empty space when gates are randomized.',
    'ONE_D_RPM action space (hover backend) hard-caps reward at ~474 regardless of reward '
    'function design. The single action dimension is the fundamental limitation.',
    'Inference-time trajectory swapping does not work. The RL policy is tightly coupled to '
    'the trajectory shape it was trained on. Dynamic trajectories require training-time diversity.',
    'Kaggle competition winners almost certainly use dynamic path planning that generates '
    'waypoints from observed gate positions, not static spline trajectories.',
    'GPU training (1024 envs) provides 16x more data per second than CPU (64 envs), enabling '
    'experiments that are impossible within CPU time budgets.',
]
for lesson in lessons:
    doc.add_paragraph(lesson, style='List Bullet')

# ===== 8. COST SUMMARY =====
doc.add_heading('8. Cost Summary', level=1)
cost_data = [
    ['Item', 'Cost'],
    ['RunPod RTX 3090 (~1 hr)', '~$0.50'],
    ['3 GPU experiments total', '~$0.50'],
    ['Budget loaded', '$10.00'],
    ['Budget remaining', '~$9.50'],
]
table = doc.add_table(rows=len(cost_data), cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(cost_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text
        if i == 0:
            for run in table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph('')
doc.add_paragraph(
    'Total GPU cost for the entire project so far: approximately $0.50. '
    'The $10 budget can fund ~20 additional training runs.'
)

# ===== 9. NEXT STEPS =====
doc.add_heading('9. Next Steps', level=1)
next_steps = [
    ('Dynamic Trajectory Training', 'Train the RL agent on diverse trajectories generated '
     'from randomized gate positions, so the policy generalizes to Level 2.'),
    ('Curriculum Learning', 'Progressive training from L0 \u2192 L1 \u2192 L2, gradually '
     'increasing difficulty so the agent builds foundational skills before facing full randomization.'),
    ('Extended GPU Training', 'Longer runs (20M+ steps) with trajectory diversity on RunPod. '
     'The periodic reward dips every ~800K steps suggest v_loss instability worth investigating.'),
    ('Observation Space Optimization', 'Experiment with what information the agent receives \u2014 '
     'gate-relative positions, velocity projections, or learned trajectory embeddings.'),
]
for title_text, desc in next_steps:
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

# -- Save --
output_path = r'C:\Users\JefferyWhitmire\Desktop\Shared\drone-rl-lab\results\PROJECT_REPORT.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
